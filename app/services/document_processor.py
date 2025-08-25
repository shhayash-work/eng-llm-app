"""
文書処理サービス
"""
import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import json
from docx import Document
try:
    import PyPDF2
except ImportError:
    try:
        import pypdf as PyPDF2
    except ImportError:
        PyPDF2 = None

from openpyxl import load_workbook

from app.models.report import DocumentReport, ReportType, AnalysisResult, AnomalyDetection
from app.services.llm_service import LLMService
from app.services.vector_store import VectorStoreService
from app.config.settings import SHAREPOINT_DOCS_DIR

logger = logging.getLogger(__name__)

if PyPDF2 is None:
    logger.warning("PyPDF2/pypdf not available. PDF reading will be disabled.")

def calculate_risk_level_enum(urgency_score: int) -> 'RiskLevel':
    """urgency_scoreから表示用リスクレベルを算出"""
    from app.models.report import RiskLevel
    if urgency_score >= 7:
        return RiskLevel.HIGH
    elif urgency_score >= 4:
        return RiskLevel.MEDIUM
    else:
        return RiskLevel.LOW

class DocumentProcessor:
    """文書処理クラス"""
    
    def __init__(self, llm_provider: Optional[str] = None, create_vector_store: bool = False):
        self.llm_service = LLMService(provider=llm_provider)
        self.vector_store = VectorStoreService(create_mode=create_vector_store)
        
    def process_directory(self, directory_path: Path) -> List[DocumentReport]:
        """ディレクトリ内の全文書を処理"""
        reports = []
        
        # サポートされているファイル拡張子
        supported_extensions = {'.txt', '.pdf', '.docx', '.xlsx'}
        
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    report = self.process_single_document(file_path)
                    if report:
                        reports.append(report)
                        logger.info(f"Processed: {file_path.name}")
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
        
        return reports
    
    def process_single_document(self, file_path: Path) -> Optional[DocumentReport]:
        """単一文書を処理（統合分析1回のみ）"""
        try:
            # ファイル内容を読み込み
            content = self._read_file_content(file_path)
            if not content:
                return None
            
            # 🤖 統合LLM分析を実行（レポートタイプ判定 + メイン分析 + 分類困難検知を1回で）
            llm_result = self.llm_service.analyze_document(content)
            if not llm_result:
                logger.error(f"統合LLM分析が失敗しました（フォールバック処理なし）: {file_path.name}")
                return None
            
            # DocumentReportオブジェクトを作成（統合分析結果を使用）
            report = self._create_report_from_unified_analysis(file_path, content, llm_result)
            
            # ベクターストアに追加
            self._add_to_vector_store(report)
            
            return report
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return None
    
    def _read_file_content(self, file_path: Path) -> str:
        """ファイル内容を読み込み"""
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.txt':
                return self._read_text_file(file_path)
            elif extension == '.pdf':
                return self._read_pdf_file(file_path)
            elif extension == '.docx':
                return self._read_docx_file(file_path)
            elif extension == '.xlsx':
                return self._read_xlsx_file(file_path)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return ""
        except Exception as e:
            logger.error(f"Failed to read file {file_path}: {e}")
            return ""
    
    def _read_text_file(self, file_path: Path) -> str:
        """テキストファイルを読み込み"""
        encodings = ['utf-8', 'shift_jis', 'euc-jp', 'iso-2022-jp']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # すべてのエンコーディングで失敗した場合
        logger.warning(f"Could not decode text file: {file_path}")
        return ""
    
    def _read_pdf_file(self, file_path: Path) -> str:
        """PDFファイルを読み込み"""
        if PyPDF2 is None:
            logger.warning(f"PDF reading not available for {file_path}")
            return "PDF読み込み機能が利用できません。PyPDF2またはpypdfをインストールしてください。"
        
        text = ""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\\n"
        except Exception as e:
            logger.error(f"PDF reading failed: {e}")
        return text
    
    def _read_docx_file(self, file_path: Path) -> str:
        """Wordファイルを読み込み（汎用的）"""
        text = ""
        try:
            doc = Document(file_path)
            
            # 段落を読み込み
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    # 一般的なデータクリーニング
                    para_text = para_text.replace('\\t', ' ')
                    para_text = ' '.join(para_text.split())  # 複数空白を1つに
                    text += para_text + "\\n"
            
            # 表を読み込み
            for table in doc.tables:
                text += "\\n表:\\n"
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            # 一般的なデータクリーニング
                            cell_text = cell_text.replace('\\n', ' ').replace('\\t', ' ')
                            cell_text = ' '.join(cell_text.split())  # 複数空白を1つに
                            if cell_text not in ['', '-', '−', '該当なし', 'なし']:
                                row_data.append(cell_text)
                    if row_data:
                        text += " | ".join(row_data) + "\\n"
                        
        except Exception as e:
            logger.error(f"DOCX reading failed: {e}")
        return text
    
    def _read_xlsx_file(self, file_path: Path) -> str:
        """Excelファイルを読み込み（汎用的）"""
        text = ""
        try:
            workbook = load_workbook(file_path, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\\nシート: {sheet_name}\\n"
                
                # 全セルの内容を順次抽出
                for row in sheet.iter_rows(values_only=True):
                    row_values = []
                    for cell_value in row:
                        if cell_value is not None:
                            # 一般的なデータクリーニング
                            clean_value = str(cell_value).strip()
                            # 不要文字の削除（改行、タブ、余分な空白）
                            clean_value = clean_value.replace('\\n', ' ').replace('\\t', ' ')
                            clean_value = ' '.join(clean_value.split())  # 複数空白を1つに
                            # 空文字でない場合のみ追加
                            if clean_value and clean_value not in ['nan', 'None', 'NULL', '#N/A']:
                                row_values.append(clean_value)
                    
                    if row_values:
                        text += " | ".join(row_values) + "\\n"
                        
        except Exception as e:
            logger.error(f"XLSX reading failed: {e}")
        return text
    
    def _create_report_from_unified_analysis(self, file_path: Path, content: str, llm_result: Dict[str, Any]) -> DocumentReport:
        """統合LLM分析結果からDocumentReportを作成"""
        from app.models.report import StatusFlag, RiskLevel
        from app.services.project_mapper import ProjectMapper
        
        # レポートタイプの設定
        report_type_str = llm_result.get('report_type', 'OTHER')
        try:
            report_type = ReportType(report_type_str)
        except ValueError:
            logger.warning(f"無効なレポートタイプ: {report_type_str}、OTHERに設定")
            report_type = ReportType.OTHER
        
        # DocumentReportオブジェクトを作成
        report = DocumentReport(
            file_path=str(file_path),
            file_name=file_path.name,
            report_type=report_type,
            content=content,
            created_at=datetime.fromtimestamp(file_path.stat().st_mtime)
        )
        
        # 🤖 統合分析結果を設定
        report.requires_human_review = llm_result.get('requires_human_review', False)
        report.analysis_confidence = llm_result.get('analysis_confidence', 0.0)
        
        # 🆕 詳細信頼度・根拠情報を設定
        report.analysis_metadata = llm_result.get('analysis_metadata', {})
        
        # 項目別信頼度詳細を収集
        confidence_details = {}
        evidence_details = {}
        
        # レポートタイプの信頼度・根拠
        confidence_details['report_type'] = llm_result.get('report_type_confidence', 0.0)
        evidence_details['report_type'] = llm_result.get('report_type_evidence', '')
        
        # プロジェクト情報の信頼度・根拠
        project_info = llm_result.get('project_info', {})
        for field in ['project_id', 'station_name', 'station_number', 'aurora_plan', 'location', 'responsible_person']:
            confidence_key = f'{field}_confidence'
            evidence_key = f'{field}_evidence'
            if confidence_key in project_info:
                confidence_details[field] = project_info[confidence_key]
            if evidence_key in project_info:
                evidence_details[field] = project_info[evidence_key]
        
        # ステータス・工程の信頼度・根拠
        confidence_details['status_flag'] = llm_result.get('status_flag_confidence', 0.0)
        evidence_details['status_flag'] = llm_result.get('status_flag_evidence', '')
        confidence_details['construction_phase'] = llm_result.get('construction_phase_confidence', 0.0)
        evidence_details['construction_phase'] = llm_result.get('construction_phase_evidence', '')
        
        # 緊急度スコアの信頼度・根拠
        confidence_details['urgency_score'] = llm_result.get('urgency_score_confidence', 0.0)
        evidence_details['urgency_score'] = llm_result.get('urgency_score_evidence', '')
        
        # 遅延理由の信頼度・根拠
        delay_reasons = llm_result.get('delay_reasons', [])
        if delay_reasons:
            for i, reason in enumerate(delay_reasons):
                confidence_details[f'delay_reason_{i}'] = reason.get('confidence', 0.0)
                evidence_details[f'delay_reason_{i}'] = reason.get('evidence', '')
        
        report.confidence_details = confidence_details
        report.evidence_details = evidence_details
        
        # 🎯 プロジェクトマッピング（直接ID + ベクター検索）
        self._apply_project_mapping(report, llm_result)
        
        # 🏷️ 新フラグ体系の適用（簡略化）
        # StatusFlag設定（LLMから直接取得）
        llm_status_flag = llm_result.get('status_flag')
        if llm_status_flag:
            if llm_status_flag == '停止':
                report.status_flag = StatusFlag.STOPPED
            elif llm_status_flag == '重大な遅延':
                report.status_flag = StatusFlag.MAJOR_DELAY
            elif llm_status_flag == '軽微な遅延':
                report.status_flag = StatusFlag.MINOR_DELAY
            elif llm_status_flag == '順調':
                report.status_flag = StatusFlag.NORMAL
            else:
                report.status_flag = StatusFlag.NORMAL
        else:
            report.status_flag = StatusFlag.NORMAL
            
        # RiskLevel設定（urgency_scoreから連動ルールで算出）
        urgency_score = llm_result.get('urgency_score', 1)
        report.urgency_score = urgency_score
        report.risk_level = calculate_risk_level_enum(urgency_score)
        
        # 🆕 報告書タイプから建設工程関連性をルールベース出力
        from app.services.report_type_mapper import ReportTypeMapper
        
        phase_mapping = ReportTypeMapper.get_phase_analysis_for_report(report_type)
        report_type_phase_mapping = llm_result.get('report_type_phase_mapping', phase_mapping.get('report_type_phase_mapping', {}))
        
        # ルールベースの期待工程と実際の出力を統合
        expected_phase = ReportTypeMapper.get_expected_phase_from_report_type(report_type)
        if report_type_phase_mapping.get('expected_primary_phase') == '不明' and expected_phase != '不明':
            report_type_phase_mapping['expected_primary_phase'] = expected_phase
            report_type_phase_mapping['mapping_confidence'] = ReportTypeMapper.get_phase_mapping(report_type).get('confidence', 0.0)
            report_type_phase_mapping['mapping_description'] = ReportTypeMapper.get_phase_mapping(report_type).get('description', '')
        
        # 報告書タイプマッピング情報を保存（統合分析用）
        report.report_type_phase_mapping = report_type_phase_mapping
        
        # 信頼度・根拠詳細に追加
        confidence_details['report_type_phase_mapping'] = report_type_phase_mapping.get('mapping_confidence', 0.0)
        evidence_details['report_type_phase_mapping'] = report_type_phase_mapping.get('mapping_description', '')
        
        # 🚧 遅延理由情報の設定（15カテゴリ体系）
        report.delay_reasons = llm_result.get('delay_reasons', [])
        
        # 🆕 LLMの抽出結果を保存（デバッグ・検証用）
        report.llm_extraction_result = llm_result.get('project_info', {})
        
        # current_status処理削除: status_flagで統一
        
        # 📋 AnalysisResult作成（簡素化）
        report.analysis_result = AnalysisResult(
            summary=llm_result.get('summary', '分析結果なし'),
            issues=llm_result.get('issues', []),
            key_points=llm_result.get('key_points', []),
            confidence=report.analysis_confidence
        )
        
        # 🚨 異常検知の統合（分析困難度として扱う）
        report.anomaly_detection = AnomalyDetection(
            is_anomaly=report.requires_human_review,
            anomaly_description=f"LLMによる分析困難度: {'要確認' if report.requires_human_review else '正常'}",
            confidence=report.analysis_confidence,
            suggested_action="手動確認を推奨" if report.requires_human_review else "自動分析完了",
            requires_human_review=report.requires_human_review,
            similar_cases=[]
        )
        
        # 🔍 人的確認フラグの設定
        self._set_review_flags(report, llm_result)
        
        return report
    
    def _set_review_flags(self, report: DocumentReport, llm_result: Dict[str, Any]):
        """人的確認フラグを設定"""
        # 報告書内容確認が必要な条件
        content_review_needed = False
        
        # 1. 遅延理由が15カテゴリに分類されない（重大問題）
        delay_reasons = llm_result.get('delay_reasons', [])
        for delay_reason in delay_reasons:
            if delay_reason.get('category') == '重大問題（要人的確認）':
                content_review_needed = True
                break
        
        # 2. 必須項目が取得できなかった
        missing_required_fields = self._check_required_fields(report, llm_result)
        if missing_required_fields:
            content_review_needed = True
            # validation_issuesに追加
            for field in missing_required_fields:
                if f"必須項目不足: {field}" not in report.validation_issues:
                    report.validation_issues.append(f"必須項目不足: {field}")
        
        # 既存のvalidation_issuesもチェック
        if report.validation_issues:
            content_review_needed = True
        
        # 3. LLMが分析困難と判定
        if report.requires_human_review:
            content_review_needed = True
        
        # 4. 分析信頼度が低い
        if report.analysis_confidence < 0.7:
            content_review_needed = True
        
        report.requires_content_review = content_review_needed
        
        # 案件との紐づけ確認が必要な条件
        mapping_review_needed = False
        
        # プロジェクトマッピングが失敗または低信頼度
        if hasattr(report, 'project_mapping_info') and report.project_mapping_info:
            method = report.project_mapping_info.get('matching_method', 'unknown')
            if method == 'vector_search':
                # ベクター検索の場合は信頼度に関わらず確認が必要
                mapping_review_needed = True
            elif method in ['mapping_failed', 'vector_search_unavailable']:
                mapping_review_needed = True
        
        report.requires_mapping_review = mapping_review_needed
    
    def _check_required_fields(self, report: DocumentReport, llm_result: Dict[str, Any]) -> List[str]:
        """必須項目をチェックして不足している項目を返す"""
        missing_fields = []
        
        # プロジェクト情報から取得
        project_info = llm_result.get('project_info', {})
        
        # 必須項目の定義（担当者名と基地局番号は除外）
        required_fields = {
            'プロジェクトID': report.project_id,
            'auRoraプラン名': project_info.get('aurora_plan'),
            '基地局名': project_info.get('station_name'),
            '住所': project_info.get('location'),
            '報告書タイプ': report.report_type,
            'ステータス': report.status_flag,
            'リスクレベル': report.risk_level,
            '緊急度スコア': getattr(report, 'urgency_score', None)
        }
        
        # 各項目をチェック
        for field_name, value in required_fields.items():
            if value is None or value == '不明' or value == '':
                missing_fields.append(field_name)
        
        # 遅延理由は遅延が発生している場合のみ必須
        if report.status_flag and report.status_flag.value in ['minor_delay', 'major_delay', 'stopped']:
            delay_reasons = getattr(report, 'delay_reasons', [])
            if not delay_reasons:
                missing_fields.append('遅延理由')
        
        return missing_fields
    
    def _add_to_vector_store(self, report: DocumentReport) -> bool:
        """ベクターストアに文書を追加"""
        metadata = {
            'file_name': report.file_name,
            'file_path': report.file_path,
            'report_type': report.report_type.value,
            'created_at': report.created_at.isoformat(),
            'risk_level': report.risk_level.value if report.risk_level else '低',
            'urgency_score': getattr(report, 'urgency_score', 1)
        }
        
        return self.vector_store.add_document(report.content, metadata)
    
    def _apply_project_mapping(self, report: DocumentReport, llm_result: Dict[str, Any]):
        """プロジェクトマッピングを適用（条件付きベクター検索）"""
        try:
            # 🚀 1. LLM直接ID抽出を最優先で試行
            project_info = llm_result.get('project_info', {})
            direct_project_id = project_info.get('project_id')
            
            if direct_project_id and direct_project_id != '不明':
                # LLMから直接IDが取得できた場合はベクター検索をスキップ
                report.project_id = direct_project_id
                report.project_mapping_info = {
                    'confidence_score': project_info.get('project_id_confidence', 0.8),
                    'matching_method': 'llm_direct',
                    'alternative_candidates': [],
                    'extracted_info': {'llm_id': direct_project_id, 'llm_confidence': project_info.get('project_id_confidence', 0.8)}
                }
                logger.info(f"LLM direct project_id mapping: {direct_project_id} for {report.file_name}")
                return
            
            # 🔍 2. 直接IDが取得できない場合のみベクター検索
            logger.info(f"Direct ID not found, using vector search for {report.file_name}")
            from app.services.project_mapper import ProjectMapper
            
            project_mapper = ProjectMapper()
            mapping_result = project_mapper.map_project(report.content, llm_result)
            
            if mapping_result.project_id:
                report.project_id = mapping_result.project_id
                logger.info(f"Vector search project_id: {mapping_result.project_id} "
                          f"(confidence: {mapping_result.confidence_score:.2f}) for {report.file_name}")
                
                # マッピング詳細情報を保存
                report.project_mapping_info = {
                    'confidence_score': mapping_result.confidence_score,
                    'matching_method': mapping_result.matching_method,
                    'alternative_candidates': mapping_result.alternative_candidates,
                    'extracted_info': mapping_result.extracted_info
                }
                
                # 低信頼度の場合は検証が必要としてマーク
                if mapping_result.confidence_score < 0.7:
                    report.has_unexpected_values = True
                    report.validation_issues.append(f"低信頼度プロジェクトマッピング: {mapping_result.confidence_score:.2f}")
                    
            else:
                report.project_id = None
                logger.warning(f"Failed to map project for {report.file_name}")
                report.has_unexpected_values = True
                report.validation_issues.append("プロジェクトマッピング失敗")
                
        except Exception as e:
            logger.error(f"プロジェクトマッピング失敗: {e}")
            report.project_id = None
            report.has_unexpected_values = True
            report.validation_issues.append(f"プロジェクトマッピングエラー: {str(e)}")
    

    
